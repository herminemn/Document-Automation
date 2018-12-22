from django.shortcuts import render, redirect
from django.views.generic import TemplateView
from django.http import HttpResponseRedirect, HttpResponse
from .forms import UploadFileForm
from .models import DocFile
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from docx import Document
import re


def simple_upload(request):
    context = {}
    if request.method == 'POST':
        uploaded_file = request.FILES['myfile']
        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        context['url'] = fs.url(filename)
    return render(request, 'uploaddocx/upload.html', context)


def doc_list(request):
    uploads = DocFile.objects.all()
    return render(request, 'uploaddocx/doc_list.html', {
        'uploads': uploads
    })


def upload_doc(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('doc_list')
    else:
        form = UploadFileForm()
    return render(request, 'uploaddocx/upload.html', {
        'form': form
    })


def modify(request, pk):
    if request.method == 'POST':
        return redirect('modify.html')
#         filename = DocFile.objects.get(pk=pk)
#         document = Document(filename)
#         variables = []
#         for paragraph in document.paragraphs:
#             match = re.findall(r"\{(.*?)\}", paragraph.text)
#             variables.append(match)
#         for table in document.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     match = re.findall(r"\{(.*?)\}", cell.text)
#                     variables.append(match)
#         document.save(filename)
#         return render(request, pk, 'uploaddocx/modify.html', {'variables': variables})
#     return render(request, pk, 'uploaddocx/modify.html')


def delete_uploaded_doc(request, pk):
    if request.method == 'POST':
        uploaded_doc = DocFile.objects.get(pk=pk)
        uploaded_doc.delete()
    return redirect('doc_list')



# def docx_words_replace(doc_obj, regex, replace):
#
#     for p in doc_obj.paragraphs:
#         if regex.search(p.text):
#             inline = p.runs
#
#             for i in range(len(inline)):
#                 if regex.search(inline[i].text):
#                     text = regex.sub(replace, inline[i].text)
#                     inline[i].text = text
#
#     for table in doc_obj.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 docx_words_replace(cell, regex, replace)
#
#
# my_regex = re.compile(r"\{(.*?)\}")
#
# def variable_input(request):
#     if request.method == 'POST' and request.FILES['myfile']:
#         myfile = request.FILES['myfile']
#         fs = FileSystemStorage()
#         filename = fs.save(myfile.name, myfile)
#         document = Document(filename)
#         variables = []
#         for paragraph in document.paragraphs:
#             match = re.findall(r"\{(.*?)\}", paragraph.text)
#             variables.append(match)
#         for table in document.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     match = re.findall(r"\{(.*?)\}", cell.text)
#                     variables.append(match)
#         document.save(filename)
#         return render(request, 'uploaddocx/upload.html', {'variables': variables})
#     return render(request, 'uploaddocx/upload.html')


# def download_docx(request):
#     document_pdf = Document()
#     document_pdf.add_heading('user_document', 0)
#     response = HttpResponse(content_type=application/vnd)

# def generate_pdf(request):
#     if request.method == 'POST':
#         try:
#             my_replace = User.objects.get
